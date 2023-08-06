import docx
from pymdocx.common.constants import WORD_NAMESPACE


def create_comment_doc(doc_file_path, doc_file_path_new):
    # 创建一个有评论的doc
    document = docx.Document(docx=doc_file_path)
    # for p in document.paragraphs:
    #     print(p)
    #     comment = p.add_comment('测试', author='Obay Daba', initials='od')
    document.save(doc_file_path_new)


def has_comment(p):
    return 1 if p.comments else 0


def get_comment(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    for p in document.paragraphs:
        for c in p.comments:
            print(c.text)


def get_comment_xml(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    comment_xml = document.comments_part.element.xml
    return comment_xml


def parse_p_comment(p_obj):
    # docx.text.comment.Comment
    comment_info_list = []
    for comment in p_obj.comments:
        comment_info_list.append({
            'comment_text': comment.text,
            'dtime': comment.element.date,
            'author': comment.element.author,
            'initials': comment.element.initials,
            'id': comment.element._id,
        })
    return comment_info_list


def add_new_comment(comment_part, author, dtime, comment_text, initials=''):
    comment = comment_part.add_comment(author, initials, dtime)
    comment._add_p(comment_text)
    return comment._id


def add_p(p1, p2, comments_part_obj, direction='next'):
    # 添加段落p及其批注
    if direction == 'next':
        p1._p.addnext(p2._p)
    elif direction == 'previous':
        p1._p.addprevious(p2._p)
    comment_list = parse_p_comment(p2)
    new_comment_id_list = []
    for comment in comment_list:
        comment_id = add_new_comment(comments_part_obj, comment['author'], comment['dtime'],
                                     comment['comment_text'],
                                     initials=comment['initials'])
        new_comment_id_list.append(comment_id)
    modify_oxml_element_comment_id(p2._p, new_comment_id_list)


def modify_oxml_element_comment_id(oxml_element_obj, new_comment_id_list):
    attr_name = '{' + WORD_NAMESPACE['w'] + '}' + 'id'
    for i, cs in enumerate(oxml_element_obj.xpath('.//w:commentRangeStart')):
        cs.attrib[attr_name] = str(new_comment_id_list[i])
    for i, ce in enumerate(oxml_element_obj.xpath('.//w:commentRangeEnd')):
        ce.attrib[attr_name] = str(new_comment_id_list[i])
    for i, cr in enumerate(oxml_element_obj.xpath('.//w:commentReference')):
        cr.attrib[attr_name] = str(new_comment_id_list[i])


def add_comment_2_p_end(p1, p2, comment_part_obj):
    comment_list = parse_p_comment(p2)
    for c in comment_list:
        p1.add_comment(c['comment_text'],
                       author=c['author'],
                       dtime=c['dtime'],
                       comment_part=comment_part_obj,
                       rangeStart=len(p1._p.getchildren()),
                       rangeEnd=len(p1._p.getchildren()))