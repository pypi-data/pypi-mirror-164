import docx
from lxml import etree
import numpy as np

from pymdocx.common.comment import has_comment
from pymdocx.common.constants import WORD_NAMESPACE
from pymdocx.common.revision import has_revision


def get_label_in_paragraph(p, label_name):
    full_label_name = 'w:' + label_name
    xml = p._p.xml
    if full_label_name in xml:
        tree = etree.fromstring(xml)
        labels = tree.xpath('//' + full_label_name, namespaces=WORD_NAMESPACE)
    else:
        labels = []
    return labels


def print_xml_node(lxml_obj):
    print(etree.tostring(lxml_obj, pretty_print=True, encoding='utf-8').decode("utf-8"))


def print_doc(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    print(document._body._element.xml)


def get_doc(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    return document


def get_element_comment_revision_matrix(element):
    comment_list = []
    revision_list = []
    for p in element.paragraphs:
        comment_list.append(has_comment(p))
        revision_list.append(has_revision(p))
    return np.array([comment_list, revision_list])


def add_p_next(p1, p2):
    p1._p.addnext(p2._p)


def _get_actual_p_index(has_add_mapping, doc_index, p_index):
    if doc_index in has_add_mapping.keys():
        has_add_mapping[doc_index] += 1
    else:
        has_add_mapping[doc_index] = 0
    return p_index - has_add_mapping[doc_index]