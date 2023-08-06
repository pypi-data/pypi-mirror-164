import docx


def get_ins_del(p):
    from pymdocx.common.utils import print_xml_node
    ins_list = get_ins_in_paragraph(p)
    for ins in ins_list:
        print_xml_node(ins)
    del_list = get_del_in_paragraph(p)
    for del_ in del_list:
        print_xml_node(del_)


def get_revision(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    for p in document.paragraphs:
        get_ins_del(p)


def has_revision_ins(p):
    # 去除自闭和标签
    has_ins = [e for e in get_ins_in_paragraph(p) if len(e) > 0]
    return 1 if has_ins else 0


def has_revision_del(p):
    has_del = [e for e in get_del_in_paragraph(p) if len(e) > 0]
    return 1 if has_del else 0


def has_revision(p):
    return has_revision_del(p) or has_revision_ins(p)


def get_ins_in_paragraph(p):
    from pymdocx.common.utils import get_label_in_paragraph
    # 获取新增修订
    return get_label_in_paragraph(p, 'ins')


def get_del_in_paragraph(p):
    from pymdocx.common.utils import get_label_in_paragraph
    # 获取删除修订
    return get_label_in_paragraph(p, 'del')


def parse_p_revision(p_obj):
    rev_list = []
    for r in p_obj.revisions:
        text = '增加：' + r.text if r.tagroot == 'ins' else '删除：' + r.text
        rev_list.append({
            'tagroot': r.tagroot,
            'text': text,
            'dtime': r.element.date,
            'author': r.element.author,
            'initials': '',
            'id': r.element._id,
        })
    return rev_list


def add_revision_2_p_end(p1, p2, comment_part_obj):
    rev_list = parse_p_revision(p2)
    for r in rev_list:
        p1.add_comment(r['text'],
                       author=r['author'],
                       dtime=r['dtime'],
                       comment_part=comment_part_obj,
                       rangeStart=len(p1._p.getchildren()),
                       rangeEnd=len(p1._p.getchildren()))


def remove_revision(p):
    for r in p.revisions:
        pp = r._rev.getparent()
        if r.tagroot == 'ins':
            pp.remove(r._rev)
        elif r.tagroot == 'del':
            new_run = p.add_run(r.text, '4')
            pp.replace(r._rev, new_run._r)